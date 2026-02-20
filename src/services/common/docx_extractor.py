"""
DOCX Text Extractor Service
Extracts text from DOCX documents (paragraphs + tables).
"""

from io import BytesIO
from typing import Any

from docx import Document

from src.config.logging_config import setup_logger

logger = setup_logger(__name__)


class DocxExtractor:
    """Extract text from DOCX documents using python-docx."""

    def extract_from_bytes(self, file_bytes: bytes, filename: str = "") -> dict[str, Any]:
        """Extract text from DOCX bytes.

        Args:
            file_bytes: Raw DOCX file bytes.
            filename: Original filename (for logging).

        Returns:
            Dict with 'text', 'paragraph_count', and 'char_count'.
        """
        try:
            doc = Document(BytesIO(file_bytes))
            parts: list[str] = []

            # Extract paragraphs
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    parts.append(text)

            # Extract table cells
            for table in doc.tables:
                for row in table.rows:
                    row_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_texts:
                        parts.append(" | ".join(row_texts))

            full_text = "\n\n".join(parts)
            logger.info(
                "Extracted %s characters from %s paragraphs (%s)",
                len(full_text),
                len(doc.paragraphs),
                filename,
            )
            return {
                "text": full_text,
                "paragraph_count": len(doc.paragraphs),
                "char_count": len(full_text),
            }
        except Exception as e:
            raise Exception(f"Failed to extract text from DOCX '{filename}': {e!s}") from e
